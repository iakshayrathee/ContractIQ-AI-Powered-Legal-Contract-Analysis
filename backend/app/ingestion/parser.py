"""
ingestion/parser.py
===================
Custom document parsing pipeline using PyMuPDF (PDF) and python-docx (DOCX).

Each parser returns a list of page dicts with this shape:
    {
        "page_number": int,   # 1-indexed
        "text":        str,   # raw text from that page (including serialised tables)
        "images":      list,  # list of raw image dicts (see _extract_page_images)
        "tables":      list,  # list of table dicts with html, text, and bbox
    }

Fix 1 — Table Serialisation:
  PyMuPDF's get_text("text") renders table cells as positional whitespace,
  which embeds poorly. After extracting paragraph text we detect tables via
  page.find_tables() and serialise each row as a pipe-delimited key→value
  sentence, e.g.:
    "Fee type: Monthly Retainer | Amount: USD 15,000/month | Due: 1st of month"
  These sentences are appended after the paragraph text so the downstream
  chunker and embedder can retrieve fee/table content accurately.

Image Extraction:
  For each page, _extract_page_images() calls page.get_images(full=True) and
  doc.extract_image(xref) to obtain raw bytes + dimensions.  Trivial images
  (too small, unsupported format, or too few bytes) are filtered out here at
  the parsing layer so the ingestion service only receives meaningful images.
  Vision description happens upstream in ingestion_service.py.
"""

import logging
import os
from pathlib import Path
from typing import Any

import structlog

logger = structlog.get_logger()

# ---------------------------------------------------------------------------
# Image filter thresholds (read from env; match .env.example defaults)
# ---------------------------------------------------------------------------

_IMAGE_MIN_WIDTH = int(os.environ.get("IMAGE_MIN_WIDTH", "100"))
_IMAGE_MIN_HEIGHT = int(os.environ.get("IMAGE_MIN_HEIGHT", "100"))

# Images whose raw bytes are below this threshold are almost certainly icons,
# stamps, or invisible page elements — skip them before the vision call.
_IMAGE_MIN_BYTES = 5 * 1024  # 5 KB

# Formats the vision model can accept as data-URI payloads.
_SUPPORTED_IMAGE_EXTS = {"png", "jpeg", "jpg"}


def parse_pdf(path: str) -> list[dict[str, Any]]:
    """
    Parse a PDF file and return page-level text + raw image data.

    Uses PyMuPDF (fitz) to extract text page by page.  Each page is returned
    as a separate dict so downstream chunkers can track page provenance.

    Args:
        path: Absolute or relative path to the PDF file.

    Returns:
        List of dicts, one per page:
            {
                "page_number": int,   # 1-indexed
                "text":        str,   # extracted text (including table rows)
                "images":      list,  # raw image dicts from _extract_page_images()
            }
        Pages with no extractable text (e.g., scanned-only) return {"text": ""}.

    Raises:
        FileNotFoundError: if the path does not exist.
        RuntimeError: if PyMuPDF fails to open the file.
    """
    import fitz  # PyMuPDF

    path_obj = Path(path)
    if not path_obj.exists():
        raise FileNotFoundError(f"PDF file not found: {path}")

    pages: list[dict[str, Any]] = []

    try:
        doc = fitz.open(str(path_obj))
    except Exception as exc:
        raise RuntimeError(f"PyMuPDF could not open '{path_obj.name}': {exc}") from exc

    try:
        for page_index in range(len(doc)):
            page = doc[page_index]
            text: str = page.get_text("text") or ""
            # Normalise whitespace: collapse consecutive blank lines
            text = "\n".join(
                line for line in text.splitlines()
            ).strip()

            # --- Fix 1: Table serialisation ---
            # PyMuPDF's positional text output garbles table cells into
            # whitespace-separated columns that embed as noise.  Detect tables
            # via find_tables() and append each row as a pipe-delimited
            # key→value sentence so the content embeds and retrieves cleanly.
            table_sentences = _serialise_page_tables(page)
            if table_sentences:
                text = text + "\n\n" + "\n".join(table_sentences)

            # --- Image extraction ---
            # Extract raw image data for meaningful images on this page.
            # Trivial images (logos, icons, decorative elements) are already
            # filtered inside _extract_page_images().  Vision description is
            # handled asynchronously in ingestion_service.py.
            images = _extract_page_images(doc, page, page_index)
            
            # --- Table extraction (HTML + text) ---
            # Extract tables as structured HTML for rendering + markdown text for embedding
            tables = _extract_tables_as_html(page)

            pages.append({
                "page_number": page_index + 1,
                "text": text,
                "images": images,
                "tables": tables,
            })
    finally:
        doc.close()

    total_images = sum(len(p["images"]) for p in pages)
    total_tables = sum(len(p.get("tables", [])) for p in pages)
    logger.info(
        "[Parser] PDF '%s' -> %d pages extracted, %d candidate images found, %d tables found.",
        path_obj.name,
        len(pages),
        total_images,
        total_tables,
    )
    return pages


def _extract_page_images(doc, page, page_index: int) -> list[dict[str, Any]]:
    """
    Extract raw image data from a single PyMuPDF page, filtering trivial images.

    For every embedded image reference on the page we:
      1. Retrieve raw bytes via doc.extract_image(xref).
      2. Skip if extension is not in the supported set (png/jpeg/jpg).
      3. Skip if width or height is below _IMAGE_MIN_WIDTH / _IMAGE_MIN_HEIGHT
         (catches bullet icons, horizontal rules, background textures).
      4. Skip if the raw byte count is below _IMAGE_MIN_BYTES (~5 KB)
         (catches transparent placeholders and invisible markers).

    All skipped images are logged at DEBUG level so they are visible in verbose
    runs without polluting INFO output.

    Args:
        doc:        Open fitz.Document object.
        page:       fitz.Page for the current page.
        page_index: Zero-based page index (used for logging only).

    Returns:
        List of dicts, one per retained image:
            {
                "img_index":  int,   # zero-based index within this page's image list
                "image_bytes": bytes, # raw image bytes
                "image_ext":  str,   # normalised extension, e.g. "png" or "jpeg"
                "width":      int,
                "height":     int,
            }
    """
    page_num = page_index  # zero-based; callers use page_index+1 for display
    retained: list[dict[str, Any]] = []

    try:
        images_on_page = page.get_images(full=True)
    except Exception as exc:
        logger.warning("[Parser] get_images() failed on page %d: %s", page_num + 1, exc)
        return retained

    for img_index, img in enumerate(images_on_page):
        xref = img[0]
        try:
            base_image = doc.extract_image(xref)
        except Exception as exc:
            logger.debug(
                "[Parser] extract_image xref=%d on page %d failed: %s",
                xref, page_num + 1, exc,
            )
            continue

        image_bytes: bytes = base_image.get("image", b"")
        image_ext: str = base_image.get("ext", "").lower()
        # Normalise "jpg" -> "jpeg" for data-URI mime type consistency
        if image_ext == "jpg":
            image_ext = "jpeg"
        width: int = base_image.get("width", 0)
        height: int = base_image.get("height", 0)

        # --- Filter 1: unsupported format ---
        if image_ext not in _SUPPORTED_IMAGE_EXTS:
            logger.debug(
                "[Parser] Skipped image %d on page %d: unsupported format '%s'",
                img_index, page_num + 1, image_ext,
            )
            continue

        # --- Filter 2: too small (icons, bullets, decorative lines) ---
        if width < _IMAGE_MIN_WIDTH or height < _IMAGE_MIN_HEIGHT:
            logger.debug(
                "[Parser] Skipped image %d on page %d: too small (%dx%d)",
                img_index, page_num + 1, width, height,
            )
            continue

        # --- Filter 3: too few bytes (logos, stamps, invisible markers) ---
        if len(image_bytes) < _IMAGE_MIN_BYTES:
            logger.debug(
                "[Parser] Skipped image %d on page %d: too few bytes (%d < %d)",
                img_index, page_num + 1, len(image_bytes), _IMAGE_MIN_BYTES,
            )
            continue

        retained.append({
            "img_index": img_index,
            "image_bytes": image_bytes,
            "image_ext": image_ext,
            "width": width,
            "height": height,
        })

    return retained


def _serialise_page_tables(page) -> list[str]:
    """
    Detect tables on a PyMuPDF page and serialise each data row as a
    pipe-delimited key→value sentence.

    Strategy:
    - Use page.find_tables() (available in PyMuPDF >= 1.23).
    - The first non-empty row is treated as the header.
    - Each subsequent row is emitted as:
        "Col1: val1 | Col2: val2 | ..."
      (empty cells are omitted to keep sentences compact).
    - Returns a list of sentence strings, one per data row.
    - Returns [] if find_tables is unavailable or no tables are found.
    """
    sentences: list[str] = []
    try:
        tables = page.find_tables()  # TableFinder object
        for table in tables.tables:
            raw = table.extract()  # list[list[str|None]]
            if not raw:
                continue

            # Determine headers: use first row, falling back to positional labels
            header_row = raw[0]
            headers = [
                str(cell).strip() if cell and str(cell).strip() else f"Column {j + 1}"
                for j, cell in enumerate(header_row)
            ]

            # Serialise data rows (skip header row)
            for row in raw[1:]:
                pairs = []
                for header, cell in zip(headers, row):
                    value = str(cell).strip() if cell is not None else ""
                    if value:  # omit empty cells
                        pairs.append(f"{header}: {value}")
                if pairs:
                    sentences.append(" | ".join(pairs))
    except AttributeError:
        # page.find_tables() not available in older PyMuPDF builds — skip silently
        pass
    except Exception as exc:
        logger.warning("[Parser] Table extraction failed on page: %s", exc)
    return sentences


def _extract_tables_as_html(page) -> list[dict[str, Any]]:
    """
    Extract tables from a PyMuPDF page as structured HTML + text representations.
    
    Returns a list of table dicts, each containing:
        {
            "html": str,         # HTML <table> markup with proper structure
            "text": str,         # Markdown-style text representation for embedding
            "bbox": list[float], # Bounding box [x0, y0, x1, y1] on page
        }
    
    Returns [] if find_tables is unavailable or no tables are found.
    """
    table_dicts: list[dict[str, Any]] = []
    try:
        tables = page.find_tables()  # TableFinder object
        for table in tables.tables:
            raw = table.extract()  # list[list[str|None]]
            if not raw:
                continue
            
            bbox = table.bbox if hasattr(table, 'bbox') else [0, 0, 0, 0]
            
            # Build HTML representation
            html_parts = ['<table>']
            
            # Header row
            if raw:
                html_parts.append('  <thead><tr>')
                for cell in raw[0]:
                    cell_text = str(cell).strip() if cell else ""
                    html_parts.append(f'    <th>{cell_text}</th>')
                html_parts.append('  </tr></thead>')
            
            # Data rows
            if len(raw) > 1:
                html_parts.append('  <tbody>')
                for row in raw[1:]:
                    html_parts.append('    <tr>')
                    for cell in row:
                        cell_text = str(cell).strip() if cell else ""
                        html_parts.append(f'      <td>{cell_text}</td>')
                    html_parts.append('    </tr>')
                html_parts.append('  </tbody>')
            
            html_parts.append('</table>')
            html = '\n'.join(html_parts)
            
            # Build text representation (markdown-style)
            text_parts = []
            if raw:
                # Header
                headers = [str(cell).strip() if cell else "" for cell in raw[0]]
                text_parts.append(" | ".join(headers))
                text_parts.append(" | ".join(["---"] * len(headers)))
                
                # Data rows
                for row in raw[1:]:
                    cells = [str(cell).strip() if cell else "" for cell in row]
                    text_parts.append(" | ".join(cells))
            
            text = "\n".join(text_parts)
            
            table_dicts.append({
                "html": html,
                "text": text,
                "bbox": list(bbox) if bbox else [0, 0, 0, 0],
            })
    
    except AttributeError:
        # page.find_tables() not available in older PyMuPDF builds
        pass
    except Exception as exc:
        logger.warning("[Parser] HTML table extraction failed on page: %s", exc)
    
    return table_dicts


def parse_docx(path: str) -> list[dict[str, Any]]:
    """
    Parse a DOCX file and return pseudo-page-level text.

    python-docx does not expose a native page concept, so we approximate
    pages by grouping paragraphs into chunks of ~40 paragraphs each
    (typical A4 page at 12pt).  This keeps the page_number field meaningful
    for downstream metadata while not requiring word-processor rendering.

    DOCX files do not support image extraction at the parser layer (images
    are embedded in the Open XML zip and would require a separate pass).
    Each page dict includes an empty "images" list for schema consistency.

    Args:
        path: Absolute or relative path to the DOCX file.

    Returns:
        List of dicts: {"page_number": int, "text": str, "images": list, "tables": list}.

    Raises:
        FileNotFoundError: if the path does not exist.
        RuntimeError: if python-docx fails to open the file.
    """
    from docx import Document as DocxDocument  # type: ignore[import]

    path_obj = Path(path)
    if not path_obj.exists():
        raise FileNotFoundError(f"DOCX file not found: {path}")

    try:
        doc = DocxDocument(str(path_obj))
    except Exception as exc:
        raise RuntimeError(f"python-docx could not open '{path_obj.name}': {exc}") from exc

    # Collect all paragraph texts (filter blanks)
    all_paragraphs: list[str] = [
        para.text.strip()
        for para in doc.paragraphs
        if para.text.strip()
    ]

    # Extract all tables from DOCX
    docx_tables = _extract_docx_tables(doc)

    if not all_paragraphs and not docx_tables:
        logger.warning("[Parser] DOCX '%s' yielded no paragraph text or tables.", path_obj.name)
        return [{"page_number": 1, "text": "", "images": [], "tables": []}]

    # Group into pseudo-pages of ~40 paragraphs
    PARAGRAPHS_PER_PAGE = 40
    pages: list[dict[str, Any]] = []
    
    # Distribute tables evenly across pseudo-pages based on position
    tables_per_page = len(docx_tables) // max(1, (len(all_paragraphs) // PARAGRAPHS_PER_PAGE + 1))
    table_idx = 0
    
    for page_idx in range(0, len(all_paragraphs), PARAGRAPHS_PER_PAGE):
        chunk_paras = all_paragraphs[page_idx: page_idx + PARAGRAPHS_PER_PAGE]
        page_num = (page_idx // PARAGRAPHS_PER_PAGE) + 1
        
        # Assign tables to this pseudo-page
        page_tables = []
        if table_idx < len(docx_tables):
            page_tables = docx_tables[table_idx:table_idx + max(1, tables_per_page)]
            table_idx += len(page_tables)
        
        pages.append({
            "page_number": page_num,
            "text": "\n".join(chunk_paras),
            "images": [],  # DOCX image extraction not yet implemented
            "tables": page_tables,
        })
    
    # If we have leftover tables, append them to the last page
    if table_idx < len(docx_tables):
        if pages:
            pages[-1]["tables"].extend(docx_tables[table_idx:])
        else:
            pages.append({
                "page_number": 1,
                "text": "",
                "images": [],
                "tables": docx_tables[table_idx:],
            })

    logger.info(
        "[Parser] DOCX '%s' -> %d pseudo-pages from %d paragraphs, %d tables extracted.",
        path_obj.name,
        len(pages),
        len(all_paragraphs),
        len(docx_tables),
    )
    return pages


def _extract_docx_tables(doc) -> list[dict[str, Any]]:
    """
    Extract all tables from a python-docx Document as HTML + text.
    
    Returns a list of table dicts, each containing:
        {
            "html": str,  # HTML <table> markup
            "text": str,  # Markdown-style text representation
            "bbox": list, # Placeholder [0, 0, 0, 0] (DOCX has no bbox)
        }
    """
    table_dicts: list[dict[str, Any]] = []
    
    for table in doc.tables:
        try:
            # Extract all rows
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(cells)
            
            if not rows:
                continue
            
            # Build HTML representation
            html_parts = ['<table>']
            
            # Header row (first row)
            if rows:
                html_parts.append('  <thead><tr>')
                for cell in rows[0]:
                    html_parts.append(f'    <th>{cell}</th>')
                html_parts.append('  </tr></thead>')
            
            # Data rows
            if len(rows) > 1:
                html_parts.append('  <tbody>')
                for row in rows[1:]:
                    html_parts.append('    <tr>')
                    for cell in row:
                        html_parts.append(f'      <td>{cell}</td>')
                    html_parts.append('    </tr>')
                html_parts.append('  </tbody>')
            
            html_parts.append('</table>')
            html = '\n'.join(html_parts)
            
            # Build text representation (markdown-style)
            text_parts = []
            if rows:
                # Header
                text_parts.append(" | ".join(rows[0]))
                text_parts.append(" | ".join(["---"] * len(rows[0])))
                
                # Data rows
                for row in rows[1:]:
                    text_parts.append(" | ".join(row))
            
            text = "\n".join(text_parts)
            
            table_dicts.append({
                "html": html,
                "text": text,
                "bbox": [0, 0, 0, 0],  # DOCX has no spatial positioning
            })
        
        except Exception as exc:
            logger.warning("[Parser] Failed to extract DOCX table: %s", exc)
            continue
    
    return table_dicts
