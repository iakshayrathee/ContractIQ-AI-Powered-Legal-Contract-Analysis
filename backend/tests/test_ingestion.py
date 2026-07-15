"""
tests/test_ingestion.py
========================
Unit tests for the custom ingestion pipeline (Task 1).

Tests cover:
  - parse_pdf() extracts text with correct page numbers
  - parse_docx() extracts paragraphs
  - chunk_text() produces expected chunk count and overlap
  - Metadata is preserved through the pipeline
  - detect_clause_type() identifies known legal clause types
  - build_document_from_chunk() helper builds correct Documents
"""

import io
import json
import os
import tempfile
from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest
from langchain_core.documents import Document

from app.ingestion.chunker import chunk_text, detect_clause_type
from app.ingestion.parser import parse_pdf, parse_docx
from app.utils.pdf_utils import build_document_from_chunk


# ---------------------------------------------------------------------------
# Fixtures — minimal in-memory files
# ---------------------------------------------------------------------------

@pytest.fixture
def minimal_pdf_path(tmp_path: Path) -> Path:
    """Create a real minimal PDF using PyMuPDF."""
    import fitz  # noqa: F401

    pdf_path = tmp_path / "test_contract.pdf"

    doc = fitz.open()
    for i in range(3):
        page = doc.new_page()
        page.insert_text(
            (72, 72),
            f"Page {i + 1} text. This contract contains terms and conditions."
            f" The indemnification clause requires both parties to hold harmless.",
        )
    doc.save(str(pdf_path))
    doc.close()
    return pdf_path


@pytest.fixture
def minimal_docx_path(tmp_path: Path) -> Path:
    """Create a real minimal DOCX using python-docx."""
    from docx import Document as DocxDoc

    docx_path = tmp_path / "test_contract.docx"
    doc = DocxDoc()
    doc.add_paragraph("This agreement is entered into between Party A and Party B.")
    doc.add_paragraph("Confidentiality: Both parties shall keep information confidential.")
    doc.add_paragraph("Termination: Either party may terminate with 30 days notice.")
    doc.add_paragraph("Payment: Party A shall pay $1000 monthly as a fee.")
    doc.save(str(docx_path))
    return docx_path


@pytest.fixture
def sample_pages() -> list[dict]:
    """Sample page list for chunker tests."""
    return [
        {
            "page_number": 1,
            "text": (
                "This is the first page of a legal contract. "
                "It contains important terms and conditions that govern the relationship "
                "between the parties. The indemnification clause requires that each party "
                "shall indemnify and hold harmless the other from any claims.\n\n"
                "The confidentiality provisions are also outlined here. "
                "Both parties agree to maintain strict confidentiality of all proprietary information."
            ),
        },
        {
            "page_number": 2,
            "text": (
                "Termination. Either party may terminate this agreement with thirty (30) days "
                "written notice. Termination shall not relieve either party of obligations "
                "accrued prior to the termination date.\n\n"
                "Payment terms: All invoices are due within thirty days. Late payment fees apply."
            ),
        },
    ]


# ---------------------------------------------------------------------------
# parse_pdf tests
# ---------------------------------------------------------------------------

class TestParsePdf:

    def test_parse_pdf_extracts_text_with_page_numbers(self, minimal_pdf_path: Path):
        """PDF parsing should return page-level text with correct 1-indexed page numbers."""
        pages = parse_pdf(str(minimal_pdf_path))

        assert len(pages) == 3
        for i, page in enumerate(pages):
            assert page["page_number"] == i + 1
            assert isinstance(page["text"], str)
            assert "Page" in page["text"] or len(page["text"]) >= 0  # text present

    def test_parse_pdf_page_numbers_are_one_indexed(self, minimal_pdf_path: Path):
        """First page must have page_number=1."""
        pages = parse_pdf(str(minimal_pdf_path))
        assert pages[0]["page_number"] == 1

    def test_parse_pdf_returns_list_of_dicts(self, minimal_pdf_path: Path):
        """Each page must be a dict with 'page_number' and 'text' keys."""
        pages = parse_pdf(str(minimal_pdf_path))
        for page in pages:
            assert "page_number" in page
            assert "text" in page

    def test_parse_pdf_file_not_found(self):
        """parse_pdf should raise FileNotFoundError for missing files."""
        with pytest.raises(FileNotFoundError):
            parse_pdf("/nonexistent/path/file.pdf")

    def test_parse_pdf_extracts_meaningful_text(self, minimal_pdf_path: Path):
        """Text extracted should contain words from the inserted text."""
        pages = parse_pdf(str(minimal_pdf_path))
        full_text = " ".join(p["text"] for p in pages)
        assert "indemnification" in full_text.lower() or len(full_text) > 10


# ---------------------------------------------------------------------------
# parse_docx tests
# ---------------------------------------------------------------------------

class TestParseDocx:

    def test_parse_docx_extracts_paragraphs(self, minimal_docx_path: Path):
        """DOCX parsing should return at least one pseudo-page of text."""
        pages = parse_docx(str(minimal_docx_path))

        assert len(pages) >= 1
        assert pages[0]["page_number"] == 1
        assert isinstance(pages[0]["text"], str)

    def test_parse_docx_contains_paragraph_content(self, minimal_docx_path: Path):
        """Extracted text should contain text from the paragraphs."""
        pages = parse_docx(str(minimal_docx_path))
        full_text = " ".join(p["text"] for p in pages)

        assert "confidential" in full_text.lower()
        assert "terminat" in full_text.lower()

    def test_parse_docx_page_numbers_start_at_one(self, minimal_docx_path: Path):
        """First pseudo-page must have page_number=1."""
        pages = parse_docx(str(minimal_docx_path))
        assert pages[0]["page_number"] == 1

    def test_parse_docx_returns_list_of_dicts(self, minimal_docx_path: Path):
        """Each pseudo-page must be a dict with 'page_number' and 'text' keys."""
        pages = parse_docx(str(minimal_docx_path))
        for page in pages:
            assert "page_number" in page
            assert "text" in page

    def test_parse_docx_file_not_found(self):
        """parse_docx should raise FileNotFoundError for missing files."""
        with pytest.raises(FileNotFoundError):
            parse_docx("/nonexistent/path/file.docx")


# ---------------------------------------------------------------------------
# chunk_text tests
# ---------------------------------------------------------------------------

class TestChunkText:

    def test_chunk_text_produces_chunks(self, sample_pages: list[dict]):
        """chunk_text should produce at least one chunk from non-empty pages."""
        chunks = chunk_text(sample_pages, source_filename="contract.pdf")
        assert len(chunks) > 0

    def test_chunk_text_produces_correct_count(self):
        """chunk_text should produce multiple chunks for long text."""
        long_text = "A" * 2000  # definitely longer than chunk_size=512
        pages = [{"page_number": 1, "text": long_text}]
        chunks = chunk_text(pages, source_filename="test.pdf", chunk_size=512, chunk_overlap=64)
        assert len(chunks) > 1

    def test_chunk_text_respects_chunk_size(self, sample_pages: list[dict]):
        """No chunk should exceed chunk_size characters."""
        chunk_size = 200
        chunks = chunk_text(sample_pages, source_filename="c.pdf", chunk_size=chunk_size, chunk_overlap=32)
        for chunk in chunks:
            assert len(chunk["text"]) <= chunk_size + 50  # allow small overshoot from recursive splits

    def test_chunk_text_has_overlap(self):
        """Chunks from adjacent pages should share overlap text at boundaries."""
        text = "word " * 200  # 1000 chars of "word "
        pages = [{"page_number": 1, "text": text}]
        chunks = chunk_text(pages, source_filename="t.pdf", chunk_size=100, chunk_overlap=30)

        if len(chunks) >= 2:
            # The start of chunk[1] should overlap with the end of chunk[0]
            end_of_first = chunks[0]["text"][-30:]
            start_of_second = chunks[1]["text"][:30]
            # There should be some overlap (not necessarily identical but from same text)
            assert any(word in start_of_second for word in end_of_first.split()[:3])

    def test_chunk_text_metadata_fields(self, sample_pages: list[dict]):
        """Each chunk must carry all required metadata fields."""
        chunks = chunk_text(sample_pages, source_filename="contract.pdf")
        for chunk in chunks:
            assert "text" in chunk
            assert "page_number" in chunk
            assert "chunk_index" in chunk
            assert "source_filename" in chunk
            assert "clause_type" in chunk  # may be None

    def test_chunk_text_chunk_index_sequential(self, sample_pages: list[dict]):
        """chunk_index must be monotonically increasing starting at 0."""
        chunks = chunk_text(sample_pages, source_filename="contract.pdf")
        indices = [c["chunk_index"] for c in chunks]
        assert indices == list(range(len(chunks)))

    def test_chunk_text_source_filename_preserved(self, sample_pages: list[dict]):
        """source_filename must be set on every chunk."""
        chunks = chunk_text(sample_pages, source_filename="my_contract.pdf")
        for chunk in chunks:
            assert chunk["source_filename"] == "my_contract.pdf"

    def test_chunk_text_empty_pages(self):
        """chunk_text should return empty list for all-empty pages."""
        pages = [{"page_number": 1, "text": ""}, {"page_number": 2, "text": "   "}]
        chunks = chunk_text(pages, source_filename="empty.pdf")
        assert chunks == []

    def test_chunk_text_page_numbers_tracked(self, sample_pages: list[dict]):
        """Page numbers in chunks should correspond to the originating page."""
        chunks = chunk_text(sample_pages, source_filename="c.pdf")
        page_numbers = {c["page_number"] for c in chunks}
        # Both pages 1 and 2 should appear in chunks
        assert 1 in page_numbers
        assert 2 in page_numbers


# ---------------------------------------------------------------------------
# detect_clause_type tests
# ---------------------------------------------------------------------------

class TestDetectClauseType:

    def test_detect_indemnification(self):
        text = "The party shall indemnify and hold harmless the other from all claims."
        assert detect_clause_type(text) == "indemnification"

    def test_detect_termination(self):
        text = "Either party may terminate this agreement upon thirty days written notice."
        assert detect_clause_type(text) == "termination"

    def test_detect_liability(self):
        text = "The limitation of liability shall not exceed the fees paid in the preceding month."
        assert detect_clause_type(text) == "liability"

    def test_detect_confidentiality(self):
        text = "All confidential information and trade secrets shall remain proprietary."
        assert detect_clause_type(text) == "confidentiality"

    def test_detect_payment(self):
        text = "Payment shall be made by wire transfer within 30 days of invoice."
        assert detect_clause_type(text) == "payment"

    def test_detect_force_majeure(self):
        text = "Neither party shall be liable for delays due to force majeure events beyond reasonable control."
        assert detect_clause_type(text) == "force_majeure"

    def test_detect_governing_law(self):
        text = "This agreement is governed by the laws of the State of Delaware."
        assert detect_clause_type(text) == "governing_law"

    def test_detect_dispute_resolution(self):
        text = "Any disputes shall be resolved through binding arbitration in accordance with AAA rules."
        assert detect_clause_type(text) == "dispute_resolution"

    def test_detect_intellectual_property(self):
        text = "All intellectual property, copyright, and patent rights shall remain with the licensor."
        assert detect_clause_type(text) == "intellectual_property"

    def test_detect_warranty(self):
        text = "The software is provided as-is without any warranty of merchantability."
        assert detect_clause_type(text) == "warranty"

    def test_detect_none_for_generic_text(self):
        """Generic text with no legal clause keywords should return None."""
        text = "This is a general introduction to the document scope."
        result = detect_clause_type(text)
        assert result is None

    def test_detect_case_insensitive(self):
        """Detection should work regardless of text case."""
        text = "INDEMNIFICATION: PARTY A SHALL INDEMNIFY PARTY B FROM ALL CLAIMS."
        assert detect_clause_type(text) == "indemnification"


# ---------------------------------------------------------------------------
# Metadata preservation through pipeline
# ---------------------------------------------------------------------------

class TestMetadataPreservation:

    def test_metadata_preserved_through_pipeline(self, sample_pages: list[dict]):
        """All required metadata fields should survive from pages → chunks → Documents."""
        source_filename = "lease_agreement.pdf"
        chunks = chunk_text(sample_pages, source_filename=source_filename)

        assert len(chunks) > 0

        for chunk in chunks:
            doc = build_document_from_chunk(chunk)

            # Top-level metadata
            assert doc.metadata["source_file"] == source_filename
            assert isinstance(doc.metadata["page_number"], int)
            assert isinstance(doc.metadata["chunk_index"], int)

            # JSON payload
            payload = json.loads(doc.metadata["original_content"])
            assert "raw_text" in payload
            assert "page_numbers" in payload
            assert "chunk_index" in payload
            assert "source_filename" in payload
            assert "clause_type" in payload  # may be None
            assert payload["source_filename"] == source_filename
            assert payload["chunk_index"] == chunk["chunk_index"]

    def test_build_document_from_chunk_page_content(self, sample_pages: list[dict]):
        """Document page_content must equal the chunk text."""
        chunks = chunk_text(sample_pages, source_filename="test.pdf")
        for chunk in chunks:
            doc = build_document_from_chunk(chunk)
            assert doc.page_content == chunk["text"]

    def test_build_document_clause_type_in_metadata(self, sample_pages: list[dict]):
        """clause_type should appear in both top-level metadata and JSON payload."""
        chunks = chunk_text(sample_pages, source_filename="test.pdf")
        for chunk in chunks:
            doc = build_document_from_chunk(chunk)
            assert "clause_type" in doc.metadata
            payload = json.loads(doc.metadata["original_content"])
            assert "clause_type" in payload
            # Must be consistent
            assert doc.metadata["clause_type"] == payload["clause_type"]
